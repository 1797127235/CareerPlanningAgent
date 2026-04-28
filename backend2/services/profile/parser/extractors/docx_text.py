"""DOCX 提取器 — 使用 python-docx 提取文本。"""
from __future__ import annotations

import io
import logging

from backend2.schemas.profile import ResumeDocument
from backend2.services.profile.parser.base import TextExtractor

logger = logging.getLogger(__name__)


class DocxTextExtractor(TextExtractor):
    """从 .docx 文件中提取文本。"""

    name = "docx"

    def supports(self, filename: str, content_type: str | None) -> bool:
        fn = filename.lower()
        return (
            fn.endswith(".docx")
            or content_type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    def extract(self, file_bytes: bytes, filename: str) -> ResumeDocument | None:
        try:
            import docx

            document = docx.Document(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in document.paragraphs if p.text]
            raw_text = "\n".join(paragraphs)
            return ResumeDocument(
                filename=filename,
                raw_text=raw_text,
                extractor=self.name,
            )
        except ImportError:
            logger.warning("python-docx 未安装，跳过 .docx 提取")
            return None
        except Exception as e:
            logger.warning("DOCX 提取失败: %s", e)
            return None
